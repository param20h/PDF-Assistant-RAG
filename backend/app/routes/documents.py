"""
Document management routes — upload, list, delete, and serve PDF files.
Background ingestion via Celery workers.
"""
import os
import sys
import uuid
import logging
import asyncio
import concurrent.futures
from datetime import datetime, timezone
from typing import Optional
from pathlib import Path
import shutil
import socket
import ipaddress
import tempfile
from urllib.parse import urlparse
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status, Query, BackgroundTasks, Request, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import select, func

from app.database import get_db
from app.exceptions import (
    ExternalServiceException,
    NotFoundException,
    ValidationException,
    AppException,
    ForbiddenException,
)
from app.models import User, Document
from app.schemas import (
    DocumentResponse,
    DocumentListResponse,
    DocumentStatusResponse,
    DocumentUpdate,
    ChunkSettings,
    UploadUrl,
)
from app.auth import get_current_user
from app.config import get_settings
from app.tasks import process_document
from app.services.document_ingestion import ingest_document
from app.services.layout_parser import AdvancedPDFParser

try:
    from crawl4ai import AsyncWebCrawler
    from crawl4ai.async_configs import BrowserConfig, CrawlerRunConfig
except ImportError as exc:
    AsyncWebCrawler = None
    BrowserConfig = None
    CrawlerRunConfig = None
    CRAWL4AI_IMPORT_ERROR = exc
else:
    CRAWL4AI_IMPORT_ERROR = None

from sqlalchemy import select
logger = logging.getLogger(__name__)
settings = get_settings()

router = APIRouter(prefix="/documents", tags=["Documents"])


ALLOWED_MIME_TYPES = settings.ALLOWED_MIME_TYPES

def _deserialize_doc(doc: Document) -> DocumentResponse:
    """Return a DocumentResponse with extracted_urls parsed from JSON string."""
    import json as _json
    response = DocumentResponse.model_validate(doc)
    if doc.extracted_urls:
        try:
            response = response.model_copy(
                update={"extracted_urls": _json.loads(doc.extracted_urls)}
            )
        except Exception:
            response = response.model_copy(update={"extracted_urls": []})
    return response

def _get_documents_query(
    db: Session,
    user_id: str,
    q: Optional[str] = None,
):
    """
    Build a filtered SQLAlchemy select query for documents belonging to a user.

    Applies an optional case-insensitive substring filter on ``original_name``.
    Does NOT apply pagination – callers are responsible for ``.limit()`` /
    ``.offset()`` so this helper stays reusable.

    Args:
        db:      Active database session.
        user_id: ID of the authenticated user whose documents to query.
        q:       Optional keyword to filter document names (case-insensitive).

    Returns:
        A SQLAlchemy ``Select`` statement ready for count or paginated execution.
    """
    base_query = (
        select(Document)
        .where(
            Document.user_id == user_id,
            Document.is_deleted.is_(False),
        )
    )

    if q and q.strip():
        pattern = f"%{q.strip()}%"
        base_query = base_query.where(
            Document.original_name.ilike(pattern)
        )

    return base_query

async def validate_upload(file: UploadFile):
    """Validate an uploaded file and save it to a temporary file.

    Checks the file extension, size (against `settings.MAX_UPLOAD_SIZE_MB`),
    MIME type via libmagic, and performs deep validation by attempting to
    parse the file (PDF with pypdf, DOCX with python-docx). On success,
    returns the path to a temporary saved file. 

    Args:
        file: The FastAPI UploadFile object to validate.

    Returns:
        str: Path to the temporary saved file that passed all validations.
    
    Raises:
        HTTPException: With status code 400 if:
            - No filename is provided.
            - The file extension is not allowed. (only .pdf or .docx)
            - The file size exceeds the maximum limit.
            - The MIME type does not match allowed types for the extension.
            - The file is corrupted or cannot be parsed. (invalid PDF/DOCX)
        HTTPException: With status code 500 if:
            - 'python-magic' dependency is missing on the server.
    """
    if not file.filename:
        raise ValidationException("No filename provided")

    ext = Path(file.filename).suffix.lower()

    # extension without leading dot in settings
    if ext.lstrip(".") not in settings.ALLOWED_EXTENSIONS:
        raise ValidationException("Only PDF, DOCX, TEXT, AND MARKDOWN files are allowed")

    # save to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        # UploadFile.file is a file-like object
        shutil.copyfileobj(file.file, tmp)
        temp_path = tmp.name

    try:
        size = Path(temp_path).stat().st_size

        if size > settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024:
            Path(temp_path).unlink(missing_ok=True)
            raise ValidationException("File too large")

        # libmagic may not be installed in all environments — import lazily
        try:
            import magic
            # make sure you have installed libmagic in your system, otherwise it will not work
        except Exception:
            Path(temp_path).unlink(missing_ok=True)
            raise ExternalServiceException("dependency", "Server missing 'python-magic' dependency")

        mime = magic.from_file(temp_path, mime=True)

        if mime not in ALLOWED_MIME_TYPES.get(ext, []):
            Path(temp_path).unlink(missing_ok=True)
            raise ValidationException(f"Invalid file type: {mime}")

        # Deep validation: try to parse the file — import parsers lazily
        try:
            if ext == ".pdf":
                from pypdf import PdfReader

                PdfReader(temp_path)
            elif ext == ".docx":
                from docx import Document as DocxDocument

                DocxDocument(temp_path)
        except Exception:
            Path(temp_path).unlink(missing_ok=True)
            raise ValidationException("Corrupted or invalid file")

        return temp_path

    finally:
        """If an exception was raised above it will propagate; caller should
        remove the temp file when appropriate. We don't unlink here because
        caller will move the file on success."""
        pass


def _crawl_in_new_loop(url: str) -> str:
    """Run the async crawler in a fresh event loop on a worker thread.
    On Windows this must be a ProactorEventLoop to support subprocesses.
    """
    if sys.platform == "win32":
        loop = asyncio.ProactorEventLoop()
    else:
        loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        async def _crawl():
            browser_config = BrowserConfig()
            run_config = CrawlerRunConfig(
                excluded_tags=['form', 'header'],

                # Content processing
                process_iframes=True,
                # remove_overlay_elements=True,

                # Cache control
                # cache_mode=CacheMode.ENABLED
            )
            async with AsyncWebCrawler(config=browser_config) as crawler:
                result = await crawler.arun(url=url, config=run_config)
                return result.markdown or ""
        return loop.run_until_complete(_crawl())
    finally:
        loop.close()

@router.post("/upload", response_model=DocumentResponse, status_code=status.HTTP_202_ACCEPTED)
async def upload_document(
    request: Request = None,
    file: UploadFile = File(...),
    chunk_size: int = Form(1000),
    chunk_overlap: int = Form(200),
    background_tasks: BackgroundTasks = None,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    # ── Validate file type ───────────────────────────
    if not file.filename:
        raise ValidationException("No filename provided")

    # ── Validate chunking params ─────────────────────
    if chunk_size < 100 or chunk_size > 2000:
        raise ValidationException("Chunk size must be between 100 and 2000")
    if chunk_overlap < 0 or chunk_overlap >= chunk_size:
        raise ValidationException("Chunk overlap must be non-negative and less than chunk_size")

    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in settings.ALLOWED_EXTENSIONS:
        raise ValidationException(
            f"File type '.{ext}' not supported. Allowed: {', '.join(settings.ALLOWED_EXTENSIONS)}",
        )

    # ── Validate and save file to disk ───────────────
    temp_path = await validate_upload(file)

    user_dir = os.path.join(settings.UPLOAD_DIR, user.id)
    os.makedirs(user_dir, exist_ok=True)

    stored_filename = f"{uuid.uuid4().hex}.{ext}"
    filepath = os.path.join(user_dir, stored_filename)

    shutil.move(temp_path, filepath)
    file_size = Path(filepath).stat().st_size

    # ── Create database record ───────────────────────
    document = Document(
        user_id=user.id,
        filename=stored_filename,
        original_name=file.filename,
        file_size=file_size,
        status="pending",
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    # ── Queue background ingestion ─────────────────
    task_id = None
    try:
        task = process_document.delay(
            document_id=document.id,
            filepath=filepath,
            original_name=file.filename,
            user_id=user.id,
        )
        task_id = task.id
    except Exception as e:
        logger.warning(f"Celery queue failed, falling back to background task: {e}")
        if background_tasks:
            background_tasks.add_task(
                ingest_document,
                document_id=document.id,
                filepath=filepath,
                original_name=file.filename,
                user_id=user.id,
            )
        task_id = f"local_{uuid.uuid4().hex}"

    return DocumentResponse.model_validate(document).model_copy(update={"task_id": task_id})

@router.post("/urlupload", status_code=status.HTTP_202_ACCEPTED)
async def upload_document_url(
        payload: UploadUrl,
        request: Request = None,
        background_tasks: BackgroundTasks = None,
        user: User = Depends(get_current_user),
        db: Session = Depends(get_db),
):
    """
    Uses crawl4ai's AsyncWebCrawler in a dedicated thread with its own
    event loop. This is required on Windows because uvicorn's default
    SelectorEventLoop does not support subprocess creation (used by
    Playwright/crawl4ai), which causes a NotImplementedError.
    On Linux (production) a plain new_event_loop() is used instead.
    """
    if CRAWL4AI_IMPORT_ERROR is not None:
        raise ExternalServiceException("crawl4ai", "URL upload is unavailable because crawl4ai is not installed")

    temp_path: Optional[str] = None
    try:
        parsed = urlparse(payload.url)
        if not all([parsed.scheme, parsed.netloc]):
            raise ValidationException("Invalid URL")

        # SSRF protection
        BLOCKED_SCHEMES = {"file", "ftp", "gopher", "dict", "smb", "ldap"}
        if parsed.scheme.lower() in BLOCKED_SCHEMES:
            raise ValidationException(f"URL scheme '{parsed.scheme}' is not allowed")
        if parsed.scheme.lower() not in ("http", "https"):
            raise ValidationException("Only http and https URLs are allowed")
        try:
            hostname = parsed.hostname
            if hostname:
                addr = socket.getaddrinfo(hostname, 80)[0][4][0]
                ip = ipaddress.ip_address(addr)
                if ip.is_private or ip.is_loopback or ip.is_link_local:
                    raise ValidationException("Internal or private URLs are not allowed")
        except (socket.gaierror, ValueError, IndexError):
            raise ValidationException("Could not resolve URL host")

        # Run in a worker thread with its own event loop to avoid
        # NotImplementedError on Windows (SelectorEventLoop can't spawn subprocesses)
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            markdown = await asyncio.get_event_loop().run_in_executor(
                pool, _crawl_in_new_loop, payload.url
            )

        if not markdown:
            raise ValidationException("No content could be extracted from the URL")


        with tempfile.NamedTemporaryFile(
            mode="w",
            suffix=".txt",
            delete=False,
            encoding="utf-8",
        ) as tmp:
            tmp.write(markdown)
            temp_path = tmp.name

        # ── Move temp file to permanent user upload directory ──
        ext = "txt"
        user_dir = os.path.join(settings.UPLOAD_DIR, user.id)
        os.makedirs(user_dir, exist_ok=True)

        stored_filename = f"{uuid.uuid4().hex}.{ext}"
        filepath = os.path.join(user_dir, stored_filename)
        shutil.move(temp_path, filepath)
        temp_path = None  # file is now at filepath; no longer a temp to clean up

        file_size = Path(filepath).stat().st_size

        # ── Derive a human-readable name from the URL ─────────
        url_path = parsed.path.rstrip("/")
        original_name = f"{parsed.netloc}{url_path or ''}.txt"

        # Bind URL crawl metadata to request state and context variables
        if request is not None:
            request.state.filename = original_name
            request.state.filesize = file_size
        from app.observability import upload_filename_var, upload_filesize_var
        upload_filename_var.set(original_name)
        upload_filesize_var.set(file_size)
        logger.info(f"URL crawler crawl completed, starting ingestion: {original_name} ({file_size} bytes)")

        # ── Create database record ─────────────────────────────
        document = Document(
            user_id=user.id,
            filename=stored_filename,
            original_name=original_name,
            file_size=file_size,
            status="pending",
        )
        db.add(document)
        db.commit()
        db.refresh(document)

        # ── Queue background ingestion ───────────────────────
        task_id = None
        try:
            task = process_document.delay(
                document_id=document.id,
                filepath=filepath,
                original_name=original_name,
                user_id=user.id,
            )
            task_id = task.id
        except Exception as e:
            logger.warning(f"Celery queue failed, falling back to background task: {e}")
            if background_tasks:
                background_tasks.add_task(
                    ingest_document,
                    document_id=document.id,
                    filepath=filepath,
                    original_name=original_name,
                    user_id=user.id,
                )
            task_id = f"local_{uuid.uuid4().hex}"

        return DocumentResponse.model_validate(document).model_copy(update={"task_id": task_id})

    except AppException:
        raise
    except ValueError:
        raise ValidationException("Invalid URL")
    except Exception as e:
        logger.error(f"URL upload error: {e}")
        raise ValidationException(f"Something went wrong with URL processing: {str(e)}")
    finally:
        '''Runs whether the request succeeded, raised an HTTPException,
        or hit an unexpected error — no temp files are ever left behind.'''
        if temp_path is not None:
            Path(temp_path).unlink(missing_ok=True)

@router.get("/trash", response_model=DocumentListResponse)
def list_trash(user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    List all soft-deleted documents for the authenticated user.
    """
    docs = db.query(Document).filter(
        Document.user_id == user.id,
        Document.is_deleted.is_(True)
    ).order_by(Document.deleted_at.desc()).all()
    
    return DocumentListResponse(
        items=[DocumentResponse.model_validate(d) for d in docs],
        total=len(docs),
        page=1,
        pages=1
    )

@router.post("/{document_id}/restore", response_model=DocumentResponse)
def restore_document(document_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    """
    Restore a soft-deleted document.
    """
    doc = db.query(Document).filter(
        Document.id == document_id, 
        Document.user_id == user.id, 
        Document.is_deleted.is_(True)
    ).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found in trash")
    
    doc.is_deleted = False
    doc.deleted_at = None
    db.commit()
    db.refresh(doc)
    
    return DocumentResponse.model_validate(doc)


@router.get("/{document_id}/status", response_model=DocumentStatusResponse)
def get_document_status(
    document_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Poll processing status for a single uploaded document.

    This endpoint lets clients refresh the upload lifecycle without fetching
    the entire document list. The returned status is one of the existing
    document states: pending, processing, ready, or failed.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise NotFoundException("Document")

    return DocumentStatusResponse.model_validate(doc)


@router.get("/", response_model=DocumentListResponse)
def list_documents(
    page: int = Query(1, ge=1, description="Page number (1-indexed)"),
    per_page: int = Query(20, ge=1, le=100, description="Results per page"),
    limit: int = Query(None, ge=1, le=100, description="Alias for per_page"),
    q: Optional[str] = Query(None, description="Filter by document name (case-insensitive)"),
    query: Optional[str] = Query(None, description="Alias for q – filter by document name"),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    List documents for the authenticated user with offset pagination and
    optional keyword search on document names.

    Pagination is controlled by ``page`` and ``per_page`` (or its alias
    ``limit``).  Search is applied via ``query`` (or its short alias ``q``),
    which performs a case-insensitive substring match on ``original_name``.

    Args:
        page:     Page number to retrieve (1-indexed). Defaults to 1.
        per_page: Number of documents per page. Defaults to 20, max 100.
        limit:    Alias for per_page – whichever is supplied takes effect.
        q:        Case-insensitive substring filter on original_name.
        query:    Alias for q – whichever is supplied takes effect.
        user:     Authenticated user injected by get_current_user.
        db:       Database session injected by get_db.

    Returns:
        DocumentListResponse with items, total, page, pages, total_pages,
        limit, and query fields.
    """
    # Allow `limit` as alias for `per_page`; `query` as alias for `q`
    effective_limit = limit if limit is not None else per_page
    effective_query = query if query is not None else q
    skip = (page - 1) * effective_limit

    # ── Build filtered query via helper ───────────────────────────────────────
    base_query = _get_documents_query(db, user.id, effective_query)

    # ── Total count (before pagination) ──────────────────────────────────────
    total = db.execute(
        select(func.count()).select_from(base_query.subquery())
    ).scalar_one()

    # ── Paginated results ─────────────────────────────────────────────────────
    docs = db.execute(
        base_query
        .order_by(Document.uploaded_at.desc())
        .limit(effective_limit)
        .offset(skip)
    ).scalars().all()

    total_pages = max(1, (total + effective_limit - 1) // effective_limit)

    return DocumentListResponse(
        items=[_deserialize_doc(d) for d in docs],
        total=total,
        page=page,
        pages=total_pages,
        total_pages=total_pages,
        limit=effective_limit,
        query=effective_query,
    )


@router.patch("/{document_id}", response_model=DocumentResponse)
def update_document(
    document_id: str,
    update: DocumentUpdate,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Update an uploaded document's metadata (name and/or summary) without
    changing its stored file or vector data.

    Both fields are optional so callers can send a partial update.  If a field
    is not present (or is null) it will be left unchanged.  To clear the
    current summary send an explicit empty string.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise NotFoundException("Document")

    if str(doc.user_id) != str(user.id):
        raise ForbiddenException("You do not have permission to update this document")

    if update.name is not None:
        doc.original_name = update.name
    if update.summary is not None:
        stripped = update.summary.strip()
        doc.summary = stripped if stripped else None

    db.commit()
    db.refresh(doc)

    return _deserialize_doc(doc)


@router.get("/{document_id}", response_model=DocumentResponse)
def get_document(
    document_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Retrieve a specific document by its ID for the authenticated user.

    Fetches the document that matches both the provided `document_id` and
    the current user's ID. If no such document exists, a 404 error is raised.

    Args:
        document_id: The unique identifier of the document to retrieve.
        user: The currently authenticated user, injected by the `get_current_user` dependency.
        db: Database session, injected by the `get_db` dependency.

    Returns:
        DocumentResponse: The document record that matches the criteria, validated against the response model
        (includes id, filename, original_name, file_size, status, etc.).

    Raises:
        HTTPException: With status code 404 if the document is not found or does not belong to the authenticated user.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise NotFoundException("Document")

    return _deserialize_doc(doc)


@router.get("/{document_id}/pdf")
def serve_pdf(
    document_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Serve the PDF file for the document viewer.

    Retrieves the document from the database to verify ownership, then
    returns the actual PDF file from disk as a downloadable response.

    Args:
        document_id: The unique identifier of the document whose PDF is to be served.
        user: The currently authenticated user, injected by the `get_current_user` dependency.
        db: Database session, injected by the `get_db` dependency.

    Returns:
        FileResponse: A FastAPI FileResponse object that streams the PDF
        file to the client with the correct media type and original filename.
    
    Raises:
        HTTPException: 404 if the document does not exist or does not belong
            to the authenticated user, or if the file is missing on disk.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise NotFoundException("Document")

    filepath = os.path.join(settings.UPLOAD_DIR, user.id, doc.filename)

    if not os.path.exists(filepath):
        raise NotFoundException("File")

    return FileResponse(
        filepath,
        media_type="application/pdf",
        filename=doc.original_name,
    )


@router.delete("/{document_id}", status_code=status.HTTP_200_OK)
def delete_document(
    document_id: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """
    Soft-delete a document so it disappears from normal document APIs.

    The underlying file, vectors, graph, and chat history are retained for a
    future recycle-bin/restore flow. Normal read/list endpoints filter deleted
    documents so accidental deletion is reversible at the database level.

    Args:
        document_id: The unique identifier of the document to delete.
        user: The currently authenticated user, injected by the `get_current_user` dependency.
        db: Database session, injected by the `get_db` dependency.

    Returns:
        dict: A JSON response containing a success message confirming the deletion of the document.

    Raises:
        HTTPException: With status code 404 if the document is not found or does not belong to the authenticated user.

    Note:
        ChromaDB deletion errors are caught and logged only; they do not
        raise an HTTP exception because the main document record is already
        removed from the database.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise NotFoundException("Document")

    doc.is_deleted = True
    doc.deleted_at = datetime.now(timezone.utc)
    db.commit()

    return {"message": f"Document '{doc.original_name}' deleted successfully"}


@router.post("/{document_id}/chunk_settings", response_model=DocumentResponse)
def update_chunk_settings(
    document_id: str,
    settings_update: ChunkSettings,
    background_tasks: BackgroundTasks = None,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Update chunking settings for a specific document.

    This endpoint allows users to update the chunk size and overlap for a document and calls _injest_document fucntion in the background to re-chunk the document with new chunk parameters.

    Args:
        document_id: The unique identifier of the document to update.
        settings_update: A ChunkSettings object containing the chunk_size and chunk_overlap values.
        background_tasks: FastAPI BackgroundTasks instance for in-process fallback execution.
        user: The currently authenticated user, injected by the `get_current_user` dependency.
        db: Database session, injected by the `get_db` dependency.

    Returns:
        DocumentResponse: The updated document record, validated against the response model.

    Raises:
        HTTPException: With status code 404 if the document is not found or does not belong to the authenticated user.
        HTTPException: With status code 400 if the provided chunk size or overlap values are invalid (e.g., chunk size less than 100, or overlap greater than or equal to chunk size).
    """
    # Validate if the document exists and belongs to the user
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise NotFoundException("Document")
    
    if settings_update.chunk_size is not None:
        if settings_update.chunk_size < 100:
            raise ValidationException("Chunk size must be at least 100")
        doc.chunk_size = settings_update.chunk_size
    if settings_update.chunk_overlap is not None:
        chunk_size_val = settings_update.chunk_size if settings_update.chunk_size is not None else (doc.chunk_size or settings.CHUNK_SIZE)
        if settings_update.chunk_overlap >= chunk_size_val:
            raise ValidationException("Chunk overlap cannot be greater than or equal to chunk size")
        doc.chunk_overlap = settings_update.chunk_overlap    

    # Refresh the document record to update the chunk settings before re-ingestion
    db.commit()
    db.refresh(doc)

    # Reset document status, chunk/page counts, summary to trigger re-ingestion with new chunk settings.
    doc.status = "pending"
    doc.chunk_count = 0
    doc.page_count = 0
    doc.summary = None
    db.commit()

    # Queue ingestion with updated chunk settings. The worker reads the new
    # settings from the document record before re-chunking.
    task_id = None
    try:
        task = process_document.delay(
            document_id=doc.id,
            filepath=os.path.join(settings.UPLOAD_DIR, user.id, doc.filename), 
            original_name=doc.original_name,
            user_id=user.id,
        )
        task_id = task.id
    except Exception as e:
        logger.warning(f"Celery queue failed, falling back to background task: {e}")
        if background_tasks:
            background_tasks.add_task(
                ingest_document,
                document_id=doc.id,
                filepath=os.path.join(settings.UPLOAD_DIR, user.id, doc.filename), 
                original_name=doc.original_name,
                user_id=user.id,
            )
        task_id = f"local_{uuid.uuid4().hex}"

    # Return the updated document record with new chunk settings
    return _deserialize_doc(doc).model_copy(update={"task_id": task_id})


@router.post("/{document_id}/retry", response_model=DocumentResponse)
def retry_document_processing(
    document_id: str,
    background_tasks: BackgroundTasks = None,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Retry processing for a failed document.

    Resets the document status back to 'pending', clears error fields,
    and re-queues the document for ingestion.
    """
    doc = db.query(Document).filter(
        Document.id == document_id,
        Document.user_id == user.id,
        Document.is_deleted.is_(False),
    ).first()

    if not doc:
        raise NotFoundException("Document")

    if doc.status != "failed":
        raise ValidationException("Only failed documents can be retried")

    doc.status = "pending"
    doc.processing_progress = 0
    doc.processing_stage = "queued"
    doc.error_message = None
    doc.last_error_traceback = None
    doc.completed_at = None
    doc.chunk_count = 0
    doc.page_count = 0
    db.commit()

    # Re-queue ingestion
    filepath = os.path.join(settings.UPLOAD_DIR, user.id, doc.filename)
    task_id = None
    try:
        task = process_document.delay(
            document_id=doc.id,
            filepath=filepath,
            original_name=doc.original_name,
            user_id=user.id,
        )
        task_id = task.id
    except Exception as e:
        logger.warning(f"Celery queue failed for retry, falling back to background task: {e}")
        if background_tasks:
            background_tasks.add_task(
                ingest_document,
                document_id=doc.id,
                filepath=filepath,
                original_name=doc.original_name,
                user_id=user.id,
            )
        task_id = f"local_{uuid.uuid4().hex}"

    return _deserialize_doc(doc).model_copy(update={"task_id": task_id})
